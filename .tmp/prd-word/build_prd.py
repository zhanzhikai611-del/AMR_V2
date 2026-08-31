from __future__ import annotations

import re
import os
import shutil
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Kai\Documents\ChatGPT\AMR\AMR_V2")
SOURCE = ROOT / "Doc" / "PRD" / "AMR V1.0.0需求文档.md"
REFERENCE = ROOT / "Doc" / "PRD" / "AMR V1.0.0需求文档.docx"
TEMPLATE = ROOT / "Doc" / "PRD" / "需求文档模板.docx"
OUTPUT = ROOT / "Doc" / "PRD" / "AMR V1.0.0需求文档-研发实施版.docx"

FONT = "Noto Sans CJK SC"
INK = "15202B"
BLUE = "0C5DCC"
MUTED = "667789"
LINE = "CFD9DF"
HEADER_FILL = "155F9E"
ALT_FILL = "F4F8FB"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_twips):
    total = sum(widths_twips)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_twips[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_twips[idx] / 1440)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_inline(paragraph, text, default_color=INK, default_size=10, bold=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|<https?://[^>]+>|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), default_size, bold, default_color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), default_size, True, default_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size - 0.5, False, "375466")
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        elif token.startswith("<"):
            set_run_font(paragraph.add_run(token[1:-1]), default_size, False, BLUE)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            set_run_font(paragraph.add_run(f"{label}（{url}）"), default_size, False, BLUE)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), default_size, bold, default_color)


def clean_cell_text(text):
    return re.sub(r"\*\*|`", "", text.strip()).replace("<br>", "\n")


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    usable = 10320
    if cols == 2:
        widths = [2300, usable - 2300]
    elif cols == 3:
        widths = [2200, 1800, usable - 4000]
    elif cols == 4:
        widths = [1450, 1500, 5650, 1720]
    else:
        base = usable // cols
        widths = [base] * cols
        widths[-1] += usable - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_fill(cell, HEADER_FILL)
            elif r_idx % 2 == 0:
                set_cell_fill(cell, ALT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.text = ""
            add_inline(p, clean_cell_text(value), "FFFFFF" if r_idx == 0 else INK, 9, r_idx == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(doc, md_path, alt_text, caption_counter):
    path = (SOURCE.parent / md_path).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(f"[截图待补充：{path.name}]"), 9, False, MUTED)
        return caption_counter
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.65))
    shape._inline.docPr.set("descr", alt_text or path.stem)
    shape._inline.docPr.set("title", alt_text or path.stem)
    return caption_counter + 1


def configure_styles(doc):
    specs = {
        "Normal": (10, False, INK, 0, 5, 1.3),
        "Title": (24, True, INK, 0, 6, 1.0),
        "Subtitle": (12, False, MUTED, 0, 5, 1.0),
        "Heading 1": (18, True, INK, 12, 8, 1.05),
        "Heading 2": (15, True, BLUE, 10, 6, 1.05),
        "Heading 3": (12, True, INK, 8, 5, 1.05),
        "Heading 4": (10.5, True, BLUE, 7, 3, 1.05),
        "Caption": (9, False, MUTED, 2, 7, 1.0),
    }
    for name, (size, bold, color, before, after, spacing) in specs.items():
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = spacing
        if name.startswith("Heading"):
            pf.keep_with_next = True


def setup_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.add_run("AMR V1.0.0需求文档  ·  "), 8.5, False, MUTED)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, 8.5, False, MUTED)


def cleanup_unused_document_images(path):
    rels_name = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(path, "r") as source:
        document_xml = source.read("word/document.xml")
        rels_root = ET.fromstring(source.read(rels_name))
        used_ids = set(re.findall(rb'r:embed="([^"]+)"', document_xml))
        removed_targets = set()
        for rel in list(rels_root):
            rel_id = rel.attrib.get("Id", "").encode()
            rel_type = rel.attrib.get("Type", "")
            if rel_type.endswith("/image") and rel_id not in used_ids:
                target = rel.attrib.get("Target", "")
                removed_targets.add("word/" + target.replace("../", ""))
                rels_root.remove(rel)
        rels_xml = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
        os.close(fd)
        Path(tmp_name).unlink(missing_ok=True)
        with zipfile.ZipFile(tmp_name, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for item in source.infolist():
                if item.filename in removed_targets:
                    continue
                data = rels_xml if item.filename == rels_name else source.read(item.filename)
                target_zip.writestr(item, data)
    Path(tmp_name).replace(path)


def build():
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    setup_footer(section)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    caption_counter = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in line.strip("|").split("|")] for line in table_lines]
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            caption_counter = add_image(doc, image_match.group(2), image_match.group(1), caption_counter)
            i += 1
            continue
        heading = re.match(r"^(#{1,5})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and first_title:
                brand = doc.add_paragraph()
                brand.paragraph_format.space_after = Pt(5)
                set_run_font(brand.add_run("AMR"), 11, True, BLUE)
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, text, INK, 24, True)
                first_title = False
            else:
                if level == 2 and re.match(r"\d+\.\s", text):
                    p_break = doc.add_paragraph()
                    p_break.add_run().add_break(WD_BREAK.PAGE)
                style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3", 5: "Heading 4"}.get(level, "Heading 4")
                p = doc.add_paragraph(style=style)
                add_inline(p, text, BLUE if style in ("Heading 2", "Heading 4") else INK, doc.styles[style].font.size.pt, True)
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Subtitle" if len(doc.paragraphs) < 5 else "Normal")
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.08)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, stripped.lstrip("> "), MUTED, 10 if len(doc.paragraphs) >= 5 else 12)
            i += 1
            continue
        list_match = re.match(r"^(\s*)(\d+\.|[-*])\s+(.+)$", raw)
        if list_match:
            marker, text = list_match.group(2), list_match.group(3)
            style = "List Number" if marker.endswith(".") else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, text)
            i += 1
            continue
        caption_match = re.match(r"^图\s*\d", stripped)
        if caption_match:
            p = doc.add_paragraph(style="Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, stripped, MUTED, 9)
            i += 1
            continue
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "!", ">", "---")) or re.match(r"^(\d+\.|[-*])\s+", nxt):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph(style="Normal")
        add_inline(p, " ".join(para_lines))

    core = doc.core_properties
    core.title = "AMR V1.0.0需求文档（研发实施版）"
    core.subject = "AMR 数字孪生、派单、行为树、资源、地图与系统设置需求"
    core.comments = "依据 AMR V1.0.0 版式与扩充后的研发实施需求生成"
    doc.save(OUTPUT)
    cleanup_unused_document_images(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
